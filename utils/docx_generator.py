"""
Generate branded Word Document (docx) for Full Campaign Package.
Brand: Bhavsar Growth Consulting
Style: Integrated Marketing Campaign Strategy & Assets
"""

import os
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

# Brand Colors
PRIMARY_BLUE = RGBColor(0x3B, 0x82, 0xF6)  # #3B82F6
ACCENT_CYAN = RGBColor(0x0E, 0xA5, 0xE9)   # #0EA5E9
DARK_BG = RGBColor(0x07, 0x0B, 0x14)       # #070B14
MUTED_TEXT = RGBColor(0x8B, 0x99, 0xAD)    # #8B99AD

def setup_branding(doc: Document):
    """Register styles for Bhavsar Growth Consulting."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Inter'
    font.size = Pt(11)
    
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Plus Jakarta Sans'
    h1.font.size = Pt(28)
    h1.font.bold = True
    h1.font.color.rgb = PRIMARY_BLUE
    
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Plus Jakarta Sans'
    h2.font.size = Pt(20)
    h2.font.bold = True
    h2.font.color.rgb = DARK_BG
    
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Plus Jakarta Sans'
    h3.font.size = Pt(16)
    h3.font.bold = True
    h3.font.color.rgb = ACCENT_CYAN

def add_cover_page(doc: Document, company_name: str):
    """Create a branded campaign cover page."""
    for _ in range(5): doc.add_paragraph()
    
    title = doc.add_paragraph("INTEGRATED MARKETING CAMPAIGN")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style = 'Heading 1'
    
    subtitle = doc.add_paragraph(f"Strategy & Content Package: {company_name}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.style = 'Heading 2'
    
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    prepared_p = doc.add_paragraph("Prepared by Bhavsar Growth Consulting Agentic System")
    prepared_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prepared_p.style = 'Normal'
    prepared_p.runs[0].font.color.rgb = MUTED_TEXT
    
    doc.add_page_break()

def markdown_to_docx(paragraph, text: str):
    """Parse simple markdown (bold, italic, links) and append to docx paragraph."""
    # This is a simplified version of the logic in the audit tool
    parts = re.split(r'(\*\*.*?\*\*)', str(text))
    for part in parts:
        is_bold = False
        if part.startswith('**') and part.endswith('**'):
            is_bold = True
            content = part[2:-2]
        else:
            content = part
            
        sub_parts = re.split(r'(\*.*?\*)', content)
        for sub_part in sub_parts:
            is_italic = False
            if sub_part.startswith('*') and sub_part.endswith('*'):
                is_italic = True
                sub_content = sub_part[1:-1]
            else:
                sub_content = sub_part
            
            if sub_content:
                run = paragraph.add_run(sub_content)
                run.bold = is_bold
                run.italic = is_italic

def add_formatted_section(doc, title, content, level=2):
    """Adds a heading and a body of text (possibly markdown)."""
    if not content: return
    doc.add_heading(title, level=level)
    
    if isinstance(content, str):
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line: continue
            
            style = 'Normal'
            if line.startswith('- ') or line.startswith('* '):
                style = 'List Bullet'
                line = line[2:]
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
                continue
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
                continue
            
            p = doc.add_paragraph(style=style)
            markdown_to_docx(p, line)
    elif isinstance(content, dict):
        for k, v in content.items():
            p = doc.add_paragraph(style='Normal')
            p.add_run(f"{k.replace('_', ' ').title()}: ").bold = True
            p.add_run(str(v))

def _build_campaign_doc(results: dict) -> Document:
    """Build the campaign document object from results."""
    doc = Document()
    setup_branding(doc)

    company_name = results.get("company_name", "Unknown Company")
    add_cover_page(doc, company_name)

    # 1. Executive Research Summary
    research = results.get("deep_research", "")
    if research:
        add_formatted_section(doc, "Deep Research & Company Identity", research)
        doc.add_page_break()

    # 2. Strategy & Framework
    strategy = results.get("strategy_framework", "")
    if strategy:
        add_formatted_section(doc, "Strategic Positioning Framework", strategy)
        doc.add_page_break()

    # 3. Campaign Brief
    brief = results.get("campaign_brief", {})
    if brief:
        add_formatted_section(doc, "Integrated Campaign Brief", brief)
        doc.add_page_break()

    # 4. PRD Compliance & Review
    feedback = results.get("reviewer_feedback", "")
    if feedback:
        add_formatted_section(doc, "PRD Compliance Audit & Strategic Review", feedback)
        doc.add_page_break()

    # 5. Content Library
    assets = results.get("generated_assets", [])
    if assets:
        doc.add_heading("Generated Content Library", level=1)
        for asset in assets:
            doc.add_heading(f"{asset.get('type')}: {asset.get('title')}", level=2)

            # Content
            add_formatted_section(doc, "Core Content", asset.get("content", ""), level=3)

            # Promo
            promo = asset.get("promotional_materials", "")
            if promo:
                add_formatted_section(doc, "Promotional Materials & Ad Copy", promo, level=3)

            doc.add_page_break()

    return doc


def generate_campaign_docx(results: dict, output_path: str):
    """Main generation function for Campaign Strategy & Assets (file path output)."""
    doc = _build_campaign_doc(results)
    doc.save(output_path)
    return output_path


def generate_campaign_docx_bytes(results: dict) -> io.BytesIO:
    """Generate campaign docx and return as BytesIO for cloud/streaming use."""
    doc = _build_campaign_doc(results)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
